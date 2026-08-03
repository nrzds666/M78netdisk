"""
M78 文档生成服务 — 独立于 RAG 服务运行。
FastAPI，端口 8001。

功能：
  POST /generate  — Markdown/JSON → docx/xlsx/md/txt/html/json/csv
  GET  /health    — 健康检查

技术栈：
  mistune     — Markdown AST 解析
  python-docx — Word 文档生成（Heading、Table、List、CodeBlock 样式映射）
  openpyxl    — Excel 文档生成（多表格 → 多 Sheet，自动推断数字类型）
  Jinja2      — 模板渲染（Phase 2）
"""

import os
import re
import uuid
from pathlib import Path
from datetime import datetime

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import mistune
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


# ── 配置 ──────────────────────────────────────────────

# 默认 temp 目录：优先用 DOC_GEN_TEMP_DIR 环境变量，否则用脚本所在目录下的 temp/
_SCRIPT_DIR = Path(__file__).resolve().parent
TEMP_DIR = Path(os.getenv("DOC_GEN_TEMP_DIR", str(_SCRIPT_DIR / "temp")))
MAX_FILENAME_LEN = 100
DEFAULT_FILENAME = "未命名文档"

app = FastAPI(title="M78 Doc Generator", version="1.0.0")


# ── 模型 ──────────────────────────────────────────────

class GenerateDocRequest(BaseModel):
    content: str           # LLM 输出的 Markdown / JSON 文本
    format: str            # md | txt | docx | xlsx | html | json | csv
    fileName: str          # 不含路径的文件名
    template: str | None = None  # 模板名（可选，Phase 2）

class GenerateDocResponse(BaseModel):
    filePath: str
    fileSize: int


# ── 工具函数 ──────────────────────────────────────────

def sanitize_filename(name: str) -> str:
    """清洗文件名：去除 Windows 非法字符、防路径穿越、截断长度。"""
    safe = re.sub(r'[/\\:*?"<>|]', '_', name)
    safe = safe.replace('..', '_')
    safe = safe.strip()
    if len(safe) > MAX_FILENAME_LEN:
        safe = safe[:MAX_FILENAME_LEN]
    if not safe:
        safe = DEFAULT_FILENAME
    return safe


def ensure_output_dir(format_ext: str) -> Path:
    """确保输出目录存在，返回按日期 + 格式分组的路径。"""
    today = datetime.now().strftime("%Y%m%d")
    out_dir = TEMP_DIR / format_ext / today
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def _to_windows_path(path: Path) -> str:
    """将 WSL 路径转为 Windows 路径（供 Java 后端使用）。"""
    path_str = str(path)
    m = re.match(r'^/mnt/([a-zA-Z])/', path_str)
    if m:
        drive = m.group(1).upper()
        rest = path_str[6:]  # strip "/mnt/X"
        return drive + ":" + rest.replace("/", chr(92))
    return path_str


def write_and_return(content_bytes: bytes, ext: str, format_ext: str) -> GenerateDocResponse:
    """写入磁盘并返回路径 + 大小。路径自动转为 Windows 格式（若在 WSL 中运行）。"""
    out_dir = ensure_output_dir(format_ext)
    file_path = out_dir / f"{uuid.uuid4().hex[:12]}.{ext}"
    file_path.write_bytes(content_bytes)
    return GenerateDocResponse(filePath=_to_windows_path(file_path), fileSize=len(content_bytes))


# ── Word (.docx) 生成 ─────────────────────────────────

def _apply_run_font(run, size_pt: int = 11, bold: bool = False):
    """设置 run 字体，解决中文字体问题。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _add_markdown_ast_to_docx(doc: Document, ast_nodes: list):
    """遍历 mistune AST，逐节点写入 python-docx。"""
    for node in ast_nodes:
        node_type = node.get('type', '')

        if node_type == 'heading':
            level = min(node.get('attrs', {}).get('level', 1), 9)
            text = _extract_children_text(node.get('children', []))
            if text.strip():
                heading = doc.add_heading(text, level=level)
                for run in heading.runs:
                    _apply_run_font(run, size_pt={1: 22, 2: 18, 3: 16, 4: 14}.get(level, 12),
                                    bold=True)

        elif node_type == 'paragraph':
            para = doc.add_paragraph()
            _render_inline_children(para, node.get('children', []))

        elif node_type == 'list':
            ordered = node.get('attrs', {}).get('ordered', False)
            for item in node.get('children', []):
                if item.get('type') == 'list_item':
                    _render_list_item(doc, item, ordered=ordered, level=0)

        elif node_type == 'block_quote':
            for child in node.get('children', []):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(_extract_children_text(child.get('children', [])))
                run.font.italic = True
                _apply_run_font(run)

        elif node_type == 'block_code':
            code_text = node.get('attrs', {}).get('info', '') or node.get('raw', '')
            if not code_text:
                code_text = _extract_children_text(node.get('children', []))
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            _apply_run_font(run, size_pt=9)

        elif node_type == 'table':
            _render_table(doc, node)

        elif node_type == 'thematic_break':
            doc.add_paragraph('─' * 60)

        elif node_type == 'block_html':
            text = node.get('raw', '')
            if text.strip():
                doc.add_paragraph(text.strip())


def _render_inline_children(para, children: list):
    """渲染行内元素（粗体、斜体、代码、链接等）。"""
    for child in children:
        ctype = child.get('type', '')
        text = child.get('raw', '') or _extract_children_text(child.get('children', []))

        if ctype == 'strong':
            run = para.add_run(text)
            _apply_run_font(run, bold=True)
        elif ctype == 'emphasis':
            run = para.add_run(text)
            _apply_run_font(run)
            run.font.italic = True
        elif ctype == 'codespan':
            run = para.add_run(text)
            run.font.name = 'Consolas'
            _apply_run_font(run, size_pt=9)
        elif ctype == 'link':
            run = para.add_run(text)
            _apply_run_font(run)
        elif ctype == 'linebreak':
            para.add_run('\n')
        elif ctype == 'softbreak':
            para.add_run(' ')
        else:
            if text:
                run = para.add_run(text)
                _apply_run_font(run)


def _render_list_item(doc, item, ordered: bool, level: int):
    """渲染列表项（支持嵌套）。"""
    prefix = f"{'  ' * level}{'1.' if ordered else '•'} "
    for child in item.get('children', []):
        ctype = child.get('type', '')
        if ctype in ('paragraph', 'block_text'):
            text = _extract_children_text(child.get('children', []))
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3 * (level + 1))
            run = para.add_run(prefix + text)
            _apply_run_font(run)
        elif child.get('type') == 'list':
            nested_ordered = child.get('attrs', {}).get('ordered', False)
            for nested_item in child.get('children', []):
                if nested_item.get('type') == 'list_item':
                    _render_list_item(doc, nested_item, ordered=nested_ordered, level=level + 1)


def _render_table(doc, table_node):
    """渲染 Markdown 表格 → python-docx Table。"""
    rows = []
    for child in table_node.get('children', []):
        if child.get('type') == 'table_head':
            # mistune table_head: children are table_cell directly (no table_row wrapper)
            cells = [_extract_children_text(c.get('children', []))
                     for c in child.get('children', []) if c.get('type') == 'table_cell']
            if cells:
                rows.append(cells)
        elif child.get('type') == 'table_body':
            for row_node in child.get('children', []):
                if row_node.get('type') == 'table_row':
                    cells = [_extract_children_text(c.get('children', []))
                             for c in row_node.get('children', [])]
                    if cells:
                        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()


def _extract_children_text(children: list) -> str:
    """递归提取子节点的纯文本。"""
    texts = []
    for child in children:
        ctype = child.get('type', '')
        if ctype == 'text':
            texts.append(child.get('raw', ''))
        elif ctype == 'block_text':
            texts.append(_extract_children_text(child.get('children', [])))
        elif ctype == 'softbreak':
            texts.append(' ')
        elif ctype == 'linebreak':
            texts.append('\n')
        elif ctype == 'strong':
            texts.append(_extract_children_text(child.get('children', [])))
        elif ctype == 'emphasis':
            texts.append(_extract_children_text(child.get('children', [])))
        elif ctype == 'codespan':
            texts.append(child.get('raw', ''))
        elif ctype == 'link':
            texts.append(_extract_children_text(child.get('children', [])))
        else:
            texts.append(child.get('raw', ''))
    return ''.join(texts)


def generate_docx(content: str) -> bytes:
    """Markdown → docx（通过 mistune AST + python-docx）。"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)

    try:
        parser = mistune.create_markdown(renderer=None, plugins=['table', 'strikethrough', 'footnotes', 'task_lists'])
        ast = parser(content)
        if isinstance(ast, list) and len(ast) > 0 and isinstance(ast[0], dict):
            _add_markdown_ast_to_docx(doc, ast)
        else:
            # 降级：纯文本
            for line in content.split('\n'):
                doc.add_paragraph(line)
    except Exception:
        # 降级：纯文本
        for line in content.split('\n'):
            doc.add_paragraph(line)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Excel (.xlsx) 生成 ─────────────────────────────────

_HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
_HEADER_FONT = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
_CELL_FONT = Font(name="微软雅黑", size=10)
_CELL_ALIGNMENT = Alignment(vertical="center")
_THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)


def _try_parse_number(s: str):
    """尝试将字符串转为 int/float，失败则返回原字符串。"""
    s = s.strip()
    if not s:
        return s
    for converter in (int, float):
        try:
            return converter(s)
        except ValueError:
            continue
    return s


def generate_xlsx(content: str) -> bytes:
    """Markdown → xlsx（提取表格节点 → openpyxl）。"""
    wb = Workbook()
    wb.remove(wb.active)  # 删除默认 Sheet

    try:
        parser = mistune.create_markdown(renderer=None, plugins=['table', 'strikethrough', 'footnotes', 'task_lists'])
        ast = parser(content)
    except Exception:
        # 降级：文本逐行写入单个 Sheet
        ws = wb.create_sheet("Sheet1")
        for i, line in enumerate(content.split('\n'), start=1):
            ws.cell(row=i, column=1, value=line)
        import io
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()

    # 提取所有表格节点
    tables = _extract_table_nodes(ast if isinstance(ast, list) else [ast])

    if not tables:
        # 无表格：纯文本
        ws = wb.create_sheet("Sheet1")
        for i, line in enumerate(content.split('\n'), start=1):
            ws.cell(row=i, column=1, value=line)
    else:
        for t_idx, table_data in enumerate(tables):
            sheet_name = f"Table{t_idx + 1}" if len(tables) > 1 else "Sheet1"
            ws = wb.create_sheet(sheet_name)

            for r_idx, row_data in enumerate(table_data, start=1):
                for c_idx, cell_value in enumerate(row_data, start=1):
                    cell = ws.cell(row=r_idx, column=c_idx)
                    parsed = _try_parse_number(cell_value)
                    cell.value = parsed
                    cell.alignment = _CELL_ALIGNMENT
                    cell.border = _THIN_BORDER

                    if r_idx == 1:
                        cell.fill = _HEADER_FILL
                        cell.font = _HEADER_FONT
                    else:
                        cell.font = _CELL_FONT

            # 自动列宽（限制最大宽度）
            for col_idx in range(1, len(table_data[0]) + 1):
                max_width = 0
                for row_data in table_data:
                    val = str(row_data[col_idx - 1]) if col_idx - 1 < len(row_data) else ""
                    max_width = max(max_width, len(val))
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_width * 1.2 + 2, 60)

    import io
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def _extract_table_nodes(ast_nodes: list) -> list[list[list[str]]]:
    """从 mistune AST 中提取所有表格节点的二维数组。"""
    tables = []
    for node in ast_nodes:
        if not isinstance(node, dict):
            continue
        if node.get('type') == 'table':
            rows = []
            for child in node.get('children', []):
                if child.get('type') == 'table_head':
                    # table_head: table_cell directly (no table_row wrapper)
                    cells = [_extract_children_text(c.get('children', []))
                             for c in child.get('children', []) if c.get('type') == 'table_cell']
                    if cells:
                        rows.append(cells)
                elif child.get('type') == 'table_body':
                    for row_node in child.get('children', []):
                        if row_node.get('type') == 'table_row':
                            cells = [_extract_children_text(c.get('children', []))
                                     for c in row_node.get('children', [])]
                            if cells:
                                rows.append(cells)
            if rows:
                tables.append(rows)
        tables.extend(_extract_table_nodes(node.get('children', [])))
    return tables


# ── 路由 ──────────────────────────────────────────────

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "service": "M78 Doc Generator",
        "tempDir": str(TEMP_DIR),
    }


@app.post("/generate", response_model=GenerateDocResponse)
async def generate_doc(req: GenerateDocRequest):
    try:
        fmt = req.format.lower().strip()
        safe_name = sanitize_filename(req.fileName)

        if fmt in ("md", "txt", "html", "json", "csv"):
            ext = fmt
            data = req.content.encode("utf-8")
            result_file = write_and_return(data, ext, fmt)

        elif fmt == "docx":
            data = generate_docx(req.content)
            result_file = write_and_return(data, "docx", "docx")

        elif fmt == "xlsx":
            data = generate_xlsx(req.content)
            result_file = write_and_return(data, "xlsx", "xlsx")

        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {fmt}")

        return result_file

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档生成失败: {str(e)}")


# ── 入口 ──────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
